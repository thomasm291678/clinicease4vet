from docx import Document
import os, json

base = r'D:\兽医\病例模版'
results = []
for fname in sorted(os.listdir(base)):
    if not fname.endswith('.docx'):
        continue
    path = os.path.join(base, fname)
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(' | '.join(cells))
    full = '\n'.join(lines)
    results.append({'file': fname, 'content': full})

for r in results:
    print('=' * 40, r['file'], '=' * 40)
    print(r['content'])
    print()
