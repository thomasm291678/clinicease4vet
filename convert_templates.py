"""提取14个docx模板为txt + 生成模板索引JSON"""
from docx import Document
import os, json, re

src_dir = r'D:\兽医\病例模版'
out_dir = r'c:\Users\HP\Documents\666666666666666666\clinicease4vet\backend\templates'
os.makedirs(out_dir, exist_ok=True)

templates_index = []
files = sorted([f for f in os.listdir(src_dir) if f.endswith('.docx')])

for fname in files:
    path = os.path.join(src_dir, fname)
    doc = Document(path)

    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(' | '.join(cells))

    full = '\n'.join(lines)

    # 从文件名提取编号和名称
    # 格式: 01-通用-小动物内科病例模版.docx
    match = re.match(r'(\d+)-(.+?)病例模版\.docx$', fname)
    num = match.group(1) if match else '00'
    name_part = match.group(2) if match else fname

    # 提取分类：心脏科、肿瘤科等
    category_match = re.match(r'.*?-(.+?)-', name_part)
    category = category_match.group(1) if category_match else '通用'

    # 简化ID
    tid = f't{num}'
    # 名称：去掉"小动物"前缀
    display_name = name_part.replace('小动物', '').strip('-').strip()

    # 保存txt
    txt_name = f'{num}_{display_name}.txt'
    txt_path = os.path.join(out_dir, txt_name)
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(full)

    templates_index.append({
        'id': tid,
        'name': display_name,
        'category': category,
        'file': txt_name,
        'size': len(full),
    })
    print(f'  [{tid}] {display_name} ({category}) -> {txt_name} ({len(full)} chars)')

# 保存索引
index_path = os.path.join(out_dir, 'index.json')
with open(index_path, 'w', encoding='utf-8') as f:
    json.dump(templates_index, f, ensure_ascii=False, indent=2)

print(f'\nDone: {len(templates_index)} templates saved to {out_dir}')
print(f'Index saved to {index_path}')
